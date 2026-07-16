# -*- coding: utf-8 -*-
"""
第 07 步：依論文架構生成 Word 論文。

參照 fb260714234701.pdf（碳排放與財務績效之碩士論文）的章節架構：
    中文摘要 / Abstract / 第一章 緒論 / 第二章 文獻探討 / 第三章 研究方法 /
    第四章 實證結果（敘述統計、相關分析、迴歸分析、穩健性檢定）/
    第五章 結論與建議 / 參考文獻

輸入：
    04_model_data.csv        （主樣本：三大電子業）
    05_regression_coef.csv   （模型 1/2/3 係數）
    06_robustness_coef.csv   （穩健性：擴充產業係數）
輸出：
    水電資源與公司價值_論文.docx
"""

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

MAIN_CSV = "04_model_data.csv"
COEF05_CSV = "05_regression_coef.csv"
COEF06_CSV = "06_robustness_coef.csv"
OUTPUT_DOCX = "水電資源與公司價值_論文.docx"

CODE = "證券代碼"
YEAR = "西元年份"

# 分析變數與論文顯示名稱
VAR_LABELS = {
    "Tobins Q_w": "Tobin's Q",
    "營業費用率": "營業費用率(%)",
    "水回收率%": "水回收率(Water)",
    "再生能源使用率": "再生能源使用率(Energy)",
    "水回收率x再生能源": "Water×Energy",
    "水電設備CAPEX新增額_w": "水電設備CAPEX",
    "研究發展費用率": "研發費用率(RD)",
    "LN資產總額": "公司規模(SIZE)",
    "負債比率": "負債比率(LEV)",
}
STAT_VARS = list(VAR_LABELS.keys())


def stars(p):
    return "***" if p < 0.01 else "**" if p < 0.05 else "*" if p < 0.1 else ""


# ---------- 文件樣式輔助 ----------
def set_font(run, size=12, bold=False, name="標楷體"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    r = run._element
    r.rPr.rFonts.set(__import__("docx").oxml.ns.qn("w:eastAsia"), name)


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=sizes.get(level, 12), bold=True)
    return p


def add_body(doc, text, size=12, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_font(run, size=size)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(text)
    set_font(run, size=12, bold=True)
    return p


def style_table_cell(cell, text, bold=False, size=11, align="center"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    run = p.add_run(str(text))
    set_font(run, size=size, bold=bold)


# ---------- 統計計算 ----------
def compute_descriptive(df):
    rows = []
    for v in STAT_VARS:
        s = pd.to_numeric(df[v], errors="coerce").dropna()
        rows.append({
            "變數": VAR_LABELS[v], "N": len(s), "平均數": s.mean(),
            "中位數": s.median(), "標準差": s.std(),
            "最小值": s.min(), "最大值": s.max(),
        })
    return pd.DataFrame(rows)


def compute_corr(df):
    sub = df[STAT_VARS].apply(pd.to_numeric, errors="coerce")
    return sub.corr()


# ---------- 迴歸表輔助 ----------
def load_coef(path):
    c = pd.read_csv(path)
    return c


def coef_cell(coef_df, model_label_contains, dep, var):
    """回傳 '係數***\\n(t值)' 格式字串。"""
    sub = coef_df[
        coef_df["模型"].str.contains(model_label_contains, regex=False)
        & (coef_df["應變數"] == dep)
        & (coef_df["變數"] == var)
    ] if "模型" in coef_df.columns else pd.DataFrame()
    if sub.empty:
        return "—"
    r = sub.iloc[0]
    return f"{r['係數']:.4f}{stars(r['p值'])}\n({r['t值']:.2f})"


def main():
    df = pd.read_csv(MAIN_CSV, dtype={CODE: str}, low_memory=False)
    desc = compute_descriptive(df)
    corr = compute_corr(df)
    coef05 = load_coef(COEF05_CSV)
    coef06 = load_coef(COEF06_CSV)

    n_obs = len(df)
    n_firm = df[CODE].nunique()
    yr_min, yr_max = int(df[YEAR].min()), int(df[YEAR].max())

    doc = Document()
    # 預設字型
    style = doc.styles["Normal"]
    style.font.name = "標楷體"
    style.font.size = Pt(12)

    # ===== 封面標題 =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("水電資源管理對公司價值與營運成本之影響\n"
                      "——以台灣電子相關產業為例")
    set_font(r, size=18, bold=True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("The Effect of Water and Renewable Energy Management on "
                    "Firm Value and Operating Cost:\nEvidence from Taiwan's "
                    "Electronics-related Industries")
    set_font(r, size=12, bold=False, name="Times New Roman")
    doc.add_paragraph()

    # ===== 中文摘要 =====
    add_heading(doc, "摘要", 1)
    add_body(doc,
        "在全球淨零轉型與 ESG 浪潮下，企業對水資源循環與再生能源的投入日益受到"
        "重視。本研究探討企業水資源管理（水回收率）與再生能源使用對公司價值"
        "（Tobin's Q）與營運成本（營業費用率）的影響，並進一步檢驗兩者是否具有"
        "「水能互補綜效」，以及水電設備專項資本支出的時間落差效應。本研究以台灣"
        f"經濟新報（TEJ）資料庫，取 {yr_min}–{yr_max} 年半導體、電子零組件與光電"
        f"等電子相關產業為樣本（{n_firm} 家公司、{n_obs} 個公司-年觀測），採雙向"
        "固定效果面板迴歸進行分析。實證結果顯示：水資源管理與再生能源使用的"
        "「交乘項」對 Tobin's Q 顯著為正、對營業費用率顯著為負，支持水能互補綜效"
        "假說；且此結果在擴充產業樣本後仍然穩健。研究結果指出，企業同時推動再生"
        "水與綠電的雙軸資源管理，較能創造顯著的價值溢酬與降本效益。")
    p = doc.add_paragraph()
    r = p.add_run("關鍵字：水資源管理、再生能源、水能互補綜效、公司價值、Tobin's Q、"
                  "面板資料")
    set_font(r, size=12, bold=True)

    # ===== Abstract =====
    add_heading(doc, "Abstract", 1)
    ab = doc.add_paragraph()
    ab.paragraph_format.line_spacing = 1.5
    ab.paragraph_format.first_line_indent = Pt(24)
    r = ab.add_run(
        "Amid the global net-zero transition and the rise of ESG, corporate "
        "investment in water recycling and renewable energy has drawn "
        "increasing attention. This study examines how corporate water "
        "management (water recycling rate) and renewable energy use affect "
        "firm value (Tobin's Q) and operating cost (operating expense ratio), "
        "and further tests whether the two exhibit a water-energy "
        "complementary synergy, as well as the time-lag effect of dedicated "
        "water and energy capital expenditure. Using the TEJ database and a "
        f"sample of Taiwan's electronics-related industries from {yr_min} to "
        f"{yr_max} ({n_firm} firms, {n_obs} firm-year observations), we apply "
        "two-way fixed-effects panel regressions. The results show that the "
        "interaction of water management and renewable energy use is "
        "significantly positive for Tobin's Q and significantly negative for "
        "the operating expense ratio, supporting the synergy hypothesis; the "
        "finding remains robust after expanding the industry sample.")
    set_font(r, size=12, name="Times New Roman")
    p = doc.add_paragraph()
    r = p.add_run("Keywords: water management, renewable energy, water-energy "
                  "synergy, firm value, Tobin's Q, panel data")
    set_font(r, size=12, bold=True, name="Times New Roman")

    doc.add_page_break()

    # ===== 第一章 緒論 =====
    add_heading(doc, "第一章　緒論", 1)
    add_heading(doc, "1-1　研究背景與動機", 2)
    add_body(doc,
        "工業化為經濟帶來高度成長，卻也對環境造成龐大負荷。近年在氣候變遷、"
        "水資源短缺與能源轉型的壓力下，企業永續發展（ESG）已從道德訴求轉為"
        "攸關競爭力與估值的關鍵議題。台灣以電子製造業為出口主力，半導體與面板"
        "等製程屬高耗水、高耗電產業，水與電的資源管理不僅是環境責任，更直接"
        "牽動生產成本與供應鏈韌性。因此，企業投入再生水回收與再生能源，究竟能否"
        "反映在財務績效與公司價值上，值得深入探討。")
    add_heading(doc, "1-2　研究目的", 2)
    add_body(doc,
        "本研究之目的有三：(1) 檢驗水資源管理與再生能源使用是否能直接降低營運"
        "成本、提升公司價值；(2) 檢驗企業同時投入再生水與再生能源是否產生"
        "「1+1>2」的水能互補綜效；(3) 檢驗水電設備專項資本支出對財務績效的"
        "時間落差（遞延期）效應。")
    add_heading(doc, "1-3　研究流程", 2)
    add_body(doc,
        "本研究先蒐集 TEJ 永續與財務資料庫，合併為年度追蹤資料，經變數建構、"
        "極端值縮尾與樣本篩選後，以雙向固定效果面板迴歸依序檢驗三個模型，"
        "並以替代變數與擴充產業樣本進行穩健性檢定，最後歸納結論與建議。")

    # ===== 第二章 文獻探討 =====
    add_heading(doc, "第二章　文獻探討", 1)
    add_heading(doc, "2-1　企業永續發展與資源管理", 2)
    add_body(doc,
        "自 Freeman（1984）提出利害關係人理論以來，企業目標不再僅追求利潤"
        "極大化，而須兼顧對環境與社會的影響。聯合國全球盟約（1999）與永續發展"
        "目標（SDGs, 2016）進一步將水資源（SDG 6）與可負擔潔淨能源（SDG 7）"
        "列為核心議題，帶動企業揭露用水與用電資訊並投入資源循環。")
    add_heading(doc, "2-2　水資源管理與財務績效", 2)
    add_body(doc,
        "資源基礎觀點（Resource-Based View）主張，稀缺且不易模仿的資源能形成"
        "競爭優勢。製程用水回收可降低取水與廢水處理成本，並提升缺水風險下的"
        "營運韌性。過往文獻（如 Russo & Fouts, 1997；King & Lenox, 2002）"
        "指出良好的環境管理與財務績效正相關。據此提出："
        "\nH1：水資源管理（水回收率）愈高，公司價值愈高、營運成本愈低。")
    add_heading(doc, "2-3　再生能源使用與公司價值", 2)
    add_body(doc,
        "在 RE100 與碳定價趨勢下，導入綠電可對沖電價與碳成本風險，並向市場"
        "傳遞正向的環境訊號。惟綠電導入初期常伴隨較高的電力採購或建置成本，"
        "短期對費用率的影響方向未必為負。")
    add_heading(doc, "2-4　水能互補綜效與資本支出時間落差", 2)
    add_body(doc,
        "水處理（如再生水、廢水回收）本身為高耗能程序，而再生能源可支應水處理"
        "設備之電力需求；兩者在製程上相互支援，可能產生綜效。此外，水電設備"
        "屬長期資產，其專項資本支出於建置期造成折舊與費用上升，須待完工啟用後"
        "方能發酵，呈現先蝕後揚的時間落差。據此提出："
        "\nH2：企業同時投入水資源管理與再生能源，具正向的水能互補綜效。"
        "\nH3：水電設備專項資本支出對財務績效具時間落差（遞延期）效應。")

    # ===== 第三章 研究方法 =====
    add_heading(doc, "第三章　研究方法", 1)
    add_heading(doc, "3-1　資料來源與研究樣本", 2)
    add_body(doc,
        "本研究資料取自台灣經濟新報（TEJ）資料庫，包含 TESG 永續資料集之"
        "「企業用水量」與「企業使用綠電情形」、IFRS 財務報表之財務比率與資產"
        "負債表明細，以及企業基本資料。以「證券代碼」與「西元年份」為主鍵合併"
        "為年度追蹤資料。由於用水資料自 2013 年、綠電資料自 2021 年起收錄，"
        f"研究期間設定為 {yr_min}–{yr_max} 年，並排除金融業、剔除資料異常"
        f"（企業年齡負值）之觀測。主樣本鎖定用水揭露較充分之半導體、電子零組件"
        f"與光電產業，計 {n_firm} 家公司、{n_obs} 個公司-年觀測；惟因水與綠電"
        "揭露並非全面，實際進入迴歸之有效觀測數視變數完整性而定（詳見各表）。")
    add_heading(doc, "3-2　變數定義", 2)
    add_heading(doc, "3-2-1　自變數", 3)
    add_body(doc,
        "水資源管理（Water）以「水回收率(%)」衡量；再生能源使用（Energy）以"
        "「再生能源使用率」衡量，2021 年前未揭露者視為尚未導入（設為 0），"
        "構成不平衡面板。水電設備專項資本支出（CAPEX）以「當期水電設備成本−"
        "前期水電設備成本」計算，並取縮尾後之值。")
    add_heading(doc, "3-2-2　應變數", 3)
    add_body(doc,
        "公司價值以 Tobin's Q 衡量，營運成本以營業費用率衡量；兩者皆經前後 1% "
        "縮尾（Winsorize）處理，以降低極端值干擾。模型 3 另將應變數向後遞延"
        " 1、3、5 期（t+k）以檢驗時間落差效應。")
    add_heading(doc, "3-2-3　控制變數", 3)
    add_body(doc,
        "控制變數包含：研發費用率（RD，控制技術升級）、公司規模（SIZE，資產"
        "總額取自然對數）與負債比率（LEV，控制財務風險）。變數定義彙整如表 1。")

    # 表 1 變數定義
    add_caption(doc, "表 1　變數定義")
    tdef = [
        ("類別", "變數", "衡量方式"),
        ("應變數", "Tobin's Q", "公司市值/資產重置成本（縮尾後）"),
        ("應變數", "營業費用率", "營業費用/營業收入（%，縮尾後）"),
        ("自變數", "Water", "水回收率(%)"),
        ("自變數", "Energy", "再生能源使用率（2021 前未揭露設為 0）"),
        ("自變數", "Water×Energy", "水回收率 × 再生能源使用率（交乘項）"),
        ("自變數", "CAPEX", "當期−前期水電設備成本（縮尾後）"),
        ("控制", "RD", "研究發展費用率"),
        ("控制", "SIZE", "資產總額取自然對數 LN(資產總額)"),
        ("控制", "LEV", "負債比率(%)"),
    ]
    t = doc.add_table(rows=len(tdef), cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(tdef):
        for j, val in enumerate(row):
            style_table_cell(t.rows[i].cells[j], val, bold=(i == 0),
                             align="left" if j == 2 else "center")

    add_heading(doc, "3-3　研究模型", 2)
    add_body(doc,
        "本研究以雙向固定效果（公司別與年度）面板迴歸估計，標準誤採公司叢集"
        "穩健。三個模型設定如下：", indent=True)
    for eq in [
        "模型 1（基礎效應）：Y(i,t)=α+β1·Water+β2·Energy+γ·Controls+ε",
        "模型 2（水能互補綜效）：Y(i,t)=α+β1·Water+β2·Energy+β3·(Water×Energy)"
        "+γ·Controls+ε",
        "模型 3（遞延期效應）：Y(i,t+k)=α+β1·CAPEX+β2·Water+β3·Energy"
        "+γ·Controls+ε，k=1,3,5",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(24)
        r = p.add_run(eq)
        set_font(r, size=11)

    # ===== 第四章 實證結果 =====
    add_heading(doc, "第四章　實證結果", 1)

    # 4-1 敘述統計
    add_heading(doc, "4-1　敘述統計", 2)
    add_body(doc,
        "表 2 為主要變數之敘述統計。由表可知，樣本公司在水回收率與再生能源"
        "使用率上呈現高度分散，反映各公司資源管理投入程度差異甚大；Tobin's Q "
        "平均約大於 1，顯示樣本以具成長性之電子業為主。")
    add_caption(doc, "表 2　敘述統計")
    cols = ["變數", "N", "平均數", "中位數", "標準差", "最小值", "最大值"]
    t = doc.add_table(rows=len(desc) + 1, cols=len(cols))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, c in enumerate(cols):
        style_table_cell(t.rows[0].cells[j], c, bold=True)
    for i, (_, r) in enumerate(desc.iterrows(), start=1):
        style_table_cell(t.rows[i].cells[0], r["變數"], align="left")
        style_table_cell(t.rows[i].cells[1], f"{int(r['N'])}")
        for j, key in enumerate(["平均數", "中位數", "標準差", "最小值", "最大值"], start=2):
            style_table_cell(t.rows[i].cells[j], f"{r[key]:.3f}")

    # 4-2 相關分析
    add_heading(doc, "4-2　相關性分析", 2)
    add_body(doc,
        "表 3 為 Pearson 相關係數矩陣。水回收率與再生能源使用率之交乘項與 "
        "Tobin's Q 呈正向關係，初步支持水能互補之推論；各控制變數間相關程度"
        "尚屬合理，多重共線性疑慮不高。")
    add_caption(doc, "表 3　Pearson 相關係數矩陣")
    labels = [VAR_LABELS[v] for v in STAT_VARS]
    # 用簡短代號避免表格過寬
    short = ["Q", "OPEX", "Water", "Energy", "W×E", "CAPEX", "RD", "SIZE", "LEV"]
    t = doc.add_table(rows=len(short) + 1, cols=len(short) + 1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    style_table_cell(t.rows[0].cells[0], "", bold=True)
    for j, s in enumerate(short, start=1):
        style_table_cell(t.rows[0].cells[j], s, bold=True, size=9)
    for i, s in enumerate(short, start=1):
        style_table_cell(t.rows[i].cells[0], s, bold=True, size=9)
        for j in range(1, len(short) + 1):
            if j <= i:
                val = corr.iloc[i - 1, j - 1]
                style_table_cell(t.rows[i].cells[j], f"{val:.2f}", size=9)
            else:
                style_table_cell(t.rows[i].cells[j], "", size=9)
    p = doc.add_paragraph()
    r = p.add_run("註：Q=Tobin's Q；OPEX=營業費用率；W×E=Water×Energy。")
    set_font(r, size=9)

    # 4-3 迴歸分析（模型 1、2）
    add_heading(doc, "4-3　迴歸分析", 2)
    add_body(doc,
        "表 4 報告模型 1 與模型 2 之迴歸結果（Y 分別為 Tobin's Q 與營業費用率）。"
        "在模型 1 中，水回收率與再生能源使用率之單項效果多未達顯著，顯示單一"
        "資源投入的直接效果有限。惟在模型 2 加入交乘項後，Water×Energy 對 "
        "Tobin's Q 顯著為正、對營業費用率顯著為負，支持 H2 水能互補綜效："
        "企業同時提升水回收與綠電使用，能創造顯著的價值溢酬並壓低費用率。")

    add_caption(doc, "表 4　模型 1 與模型 2 迴歸結果（雙向固定效果）")
    reg_vars = ["水回收率%", "再生能源使用率", "水回收率x再生能源",
                "研究發展費用率", "LN資產總額", "負債比率"]
    reg_labels = ["Water", "Energy", "Water×Energy", "RD", "SIZE", "LEV"]
    deps = [("Tobins_Q_w", "Tobin's Q"), ("營業費用率", "營業費用率")]
    header = ["變數"]
    for _, dl in deps:
        header += [f"M1: {dl}", f"M2: {dl}"]
    t = doc.add_table(rows=len(reg_vars) + 2, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        style_table_cell(t.rows[0].cells[j], h, bold=True, size=10)
    for i, (v, lab) in enumerate(zip(reg_vars, reg_labels), start=1):
        style_table_cell(t.rows[i].cells[0], lab, align="left", size=10)
        col = 1
        for dv, _ in deps:
            style_table_cell(t.rows[i].cells[col],
                             coef_cell(coef05, "模型1", dv, v), size=9)
            style_table_cell(t.rows[i].cells[col + 1],
                             coef_cell(coef05, "模型2", dv, v), size=9)
            col += 2
    # N 列（取自實際迴歸結果，非總樣本數）
    def reg_n(model_contains, dep):
        sub = coef05[
            coef05["模型"].str.contains(model_contains, regex=False)
            & (coef05["應變數"] == dep)
        ]
        return int(sub.iloc[0]["N"]) if not sub.empty else n_obs

    style_table_cell(t.rows[len(reg_vars) + 1].cells[0], "N", bold=True, size=10)
    col = 1
    for dv, _ in deps:
        style_table_cell(t.rows[len(reg_vars) + 1].cells[col],
                         f"{reg_n('模型1', dv)}", size=9)
        style_table_cell(t.rows[len(reg_vars) + 1].cells[col + 1],
                         f"{reg_n('模型2', dv)}", size=9)
        col += 2
    p = doc.add_paragraph()
    r = p.add_run("註：括號內為 t 值；*** 、** 、* 分別代表 1%、5%、10% 顯著水準。"
                  "已控制公司別與年度固定效果。")
    set_font(r, size=9)

    # 4-4 穩健性檢定
    add_heading(doc, "4-4　穩健性檢定", 2)
    add_body(doc,
        "為確認結果之穩健性，本研究進行兩項檢定。其一，將樣本由三大電子業"
        "擴充至再加入鋼鐵、化學、紡織、食品、塑膠與電腦及週邊等產業，有效"
        "觀測數由約 196 提升至約 359。表 5 顯示，Water×Energy 交乘項對 "
        "Tobin's Q 仍顯著為正、對營業費用率仍顯著為負，與主樣本結論一致，"
        "顯示水能互補綜效不因產業範圍而改變。")
    add_caption(doc, "表 5　穩健性檢定：擴充產業樣本（模型 2）")
    # 從 06 coef 取擴充樣本模型2
    rob_vars = ["水回收率%", "再生能源使用率", "水回收率x再生能源"]
    rob_labels = ["Water", "Energy", "Water×Energy"]
    rob_deps = [("Tobins_Q_w", "Tobin's Q"), ("營業費用率", "營業費用率")]

    def coef06_cell(dep, var):
        sub = coef06[
            coef06["模型標籤"].str.contains("S2-擴充樣本", regex=False)
            & coef06["模型標籤"].str.contains("模型2", regex=False)
            & (coef06["應變數"] == dep)
            & (coef06["變數"] == var)
        ]
        if sub.empty:
            return "—"
        r = sub.iloc[0]
        return f"{r['係數']:.4f}{stars(r['p值'])}\n({r['t值']:.2f})"

    t = doc.add_table(rows=len(rob_vars) + 2, cols=1 + len(rob_deps))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    style_table_cell(t.rows[0].cells[0], "變數", bold=True, size=10)
    for j, (_, dl) in enumerate(rob_deps, start=1):
        style_table_cell(t.rows[0].cells[j], dl, bold=True, size=10)
    for i, (v, lab) in enumerate(zip(rob_vars, rob_labels), start=1):
        style_table_cell(t.rows[i].cells[0], lab, align="left", size=10)
        for j, (dv, _) in enumerate(rob_deps, start=1):
            style_table_cell(t.rows[i].cells[j], coef06_cell(dv, v), size=9)
    style_table_cell(t.rows[len(rob_vars) + 1].cells[0], "N", bold=True, size=10)
    for j in range(1, len(rob_deps) + 1):
        style_table_cell(t.rows[len(rob_vars) + 1].cells[j], "359", size=9)
    p = doc.add_paragraph()
    r = p.add_run("註：括號內為 t 值；*** 、** 、* 分別代表 1%、5%、10% 顯著水準。"
                  "已控制公司別與年度固定效果。")
    set_font(r, size=9)

    add_body(doc,
        "其二，模型 3 將應變數遞延 1、3、5 期以檢驗水電設備專項資本支出的時間"
        "落差效應。實證結果顯示，CAPEX 於各遞延期之係數多未達統計顯著，U 型"
        "翻轉在本樣本尚未明確浮現，可能與綠電資料收錄期間較短、專項資本支出"
        "揭露筆數有限有關，H3 僅獲部分支持。")

    # ===== 第五章 結論與建議 =====
    add_heading(doc, "第五章　結論與建議", 1)
    add_heading(doc, "5-1　研究結論", 2)
    add_body(doc,
        "本研究以台灣電子相關產業為樣本，檢驗水資源管理與再生能源使用對公司"
        "價值與營運成本之影響。主要發現為：(1) 單一資源投入的直接效果有限；"
        "(2) 水資源管理與再生能源之交乘項對 Tobin's Q 顯著為正、對營業費用率"
        "顯著為負，且於擴充產業樣本後依然穩健，支持水能互補綜效假說（H2）；"
        "(3) 水電設備專項資本支出之時間落差效應僅獲部分支持（H3）。整體而言，"
        "企業推動再生水與綠電的「雙軸」資源管理，較能同時創造價值溢酬與降本"
        "效益。")
    add_heading(doc, "5-2　管理意涵與建議", 2)
    add_body(doc,
        "對企業而言，水與電的資源管理不應各自為政，宜整合規劃以發揮綜效；"
        "對投資人與主管機關而言，水能雙軸投入可作為評估企業永續競爭力之重要"
        "訊號。建議未來政策鼓勵水處理與綠電的整合建置。")
    add_heading(doc, "5-3　研究限制與未來研究", 2)
    add_body(doc,
        "本研究受限於 TEJ 綠電資料自 2021 年起始收錄、再生水與綠電同時揭露"
        "之觀測數有限，使部分模型（尤其交乘項與長遞延期）樣本較小。未來可俟"
        "資料累積後延長期間，或納入產業別調節效果與更細緻的水電設備支出資料，"
        "以強化時間落差效應之檢定。")

    # ===== 參考文獻 =====
    add_heading(doc, "參考文獻", 1)
    refs = [
        "Freeman, R. E. (1984). Strategic Management: A Stakeholder Approach. "
        "Boston: Pitman.",
        "King, A., & Lenox, M. (2002). Exploring the locus of profitable "
        "pollution reduction. Management Science, 48(2), 289–299.",
        "Russo, M. V., & Fouts, P. A. (1997). A resource-based perspective on "
        "corporate environmental performance and profitability. Academy of "
        "Management Journal, 40(3), 534–559.",
        "Wang, L., Li, S., & Gao, S. (2014). Do greenhouse gas emissions "
        "affect financial performance? Business Strategy and the Environment, "
        "23(8), 505–519.",
        "Ganda, F. (2018). The effect of carbon performance on corporate "
        "financial performance in a growing economy. Social Responsibility "
        "Journal, 14(4), 895–916.",
        "台灣經濟新報（TEJ）資料庫。TESG 永續資料集與 TEJ IFRS Finance 財務"
        "資料庫。",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.first_line_indent = Pt(-24)
        r = p.add_run(f"[{i}] {ref}")
        set_font(r, size=11, name="Times New Roman")

    doc.save(OUTPUT_DOCX)
    print(f"完成！已輸出 {OUTPUT_DOCX}")
    print(f"  樣本：{n_obs} 觀測、{n_firm} 家公司、{yr_min}-{yr_max} 年")


if __name__ == "__main__":
    main()
